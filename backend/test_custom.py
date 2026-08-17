import fitz
import docx
import arabic_reshaper
from bidi.algorithm import get_display
from docx.enum.text import WD_ALIGN_PARAGRAPH

def is_arabic(text):
    return any('\u0600' <= c <= '\u06FF' for c in text)

def convert_pdf_custom(pdf_path, docx_path):
    pdf = fitz.open(pdf_path)
    doc = docx.Document()
    
    for page in pdf:
        blocks = page.get_text("blocks")
        # Sort by Y position primarily, then X position
        blocks.sort(key=lambda b: (b[1], b[0]))
        
        for b in blocks:
            if b[6] != 0: # 0 means text block, 1 means image block
                continue
                
            text = b[4].strip()
            if not text:
                continue
                
            # Clean up PyMuPDF's manual line breaks within a block to allow Word to flow text naturally
            text = text.replace('\n', ' ')
            
            p = doc.add_paragraph()
            
            if is_arabic(text):
                # Shape and Bidi the ENTIRE block (paragraph)
                reshaped = arabic_reshaper.reshape(text)
                bidi_text = get_display(reshaped)
                
                run = p.add_run(bidi_text)
                run.font.name = 'Arial'
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                run = p.add_run(text)
                run.font.name = 'Arial'
                
        # Add page break
        doc.add_page_break()
        
    doc.save(docx_path)
    print("Done custom conversion:", docx_path)

if __name__ == "__main__":
    import sys
    convert_pdf_custom(sys.argv[1], sys.argv[2])
