import sys
import docx
import arabic_reshaper
from bidi.algorithm import get_display
from docx.enum.text import WD_ALIGN_PARAGRAPH

def is_arabic(text):
    return any('\u0600' <= c <= '\u06FF' for c in text)

def fix_docx_arabic(docx_path, out_path):
    doc = docx.Document(docx_path)
    
    for p in doc.paragraphs:
        if not p.text:
            continue
            
        if is_arabic(p.text):
            # Apply shaping and bidi to the whole paragraph text, but this destroys run formatting!
            # The user's code suggests:
            reshaped = arabic_reshaper.reshape(p.text)
            bidi_text = get_display(reshaped)
            
            # Since pdf2docx creates many runs, if we just replace the whole text, we lose formatting (bold, colors, links).
            # But pdf2docx usually extracts words into separate runs.
            # Let's try doing it on the paragraph level first and see if it looks right.
            # Wait, if we replace the entire paragraph text in a single run, it loses styling.
            # If we do it run by run, the visual order across the paragraph won't be correct because 'get_display' needs the full string context to order words properly.
            # Let's do it run by run first, and if that fails, paragraph level.
            # Actually, `pdf2docx` puts individual words in runs sometimes.
            # Let's just clear the paragraph and put the new text in a single run for testing.
            
            # Save paragraph formatting
            # We will just replace all runs with a single run for now to ensure correct ordering
            # because bidi algorithm must operate on the full paragraph text to order words correctly.
            
            p.clear()
            run = p.add_run(bidi_text)
            
            # Try to preserve font if possible (defaulting to Arial)
            run.font.name = 'Arial'
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Also check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if not p.text:
                        continue
                    if is_arabic(p.text):
                        reshaped = arabic_reshaper.reshape(p.text)
                        bidi_text = get_display(reshaped)
                        p.clear()
                        run = p.add_run(bidi_text)
                        run.font.name = 'Arial'
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.save(out_path)
    print("Done fixing:", out_path)

if __name__ == "__main__":
    if len(sys.argv) > 2:
        fix_docx_arabic(sys.argv[1], sys.argv[2])
