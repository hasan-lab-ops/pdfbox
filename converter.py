import pdfplumber
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import arabic_reshaper
from bidi.algorithm import get_display

def has_arabic(text):
    return any('\u0600' <= c <= '\u06FF' for c in text)

def fix_arabic(text):
    reshaped = arabic_reshaper.reshape(text)
    return get_display(reshaped)

def group_lines(words, y_threshold=5):
    lines = []
    for w in words:
        placed = False
        for line in lines:
            if abs(line[0]['top'] - w['top']) < y_threshold:
                line.append(w)
                placed = True
                break
        if not placed:
            lines.append([w])
    return lines

def sort_line(words):
    text = " ".join(w['text'] for w in words)
    if has_arabic(text):
        return sorted(words, key=lambda w: w['x0'], reverse=True)
    else:
        return sorted(words, key=lambda w: w['x0'])

def pdf_to_docx(input_path, output_path):
    doc = Document()

    with pdfplumber.open(input_path) as pdf:
        for page in pdf.pages:
            words = page.extract_words()
            lines = group_lines(words)

            for line_words in lines:
                sorted_words = sort_line(line_words)
                line_text = " ".join(w['text'] for w in sorted_words)

                line_text = line_text.encode('utf-8', 'ignore').decode('utf-8')

                if has_arabic(line_text):
                    line_text = fix_arabic(line_text)

                p = doc.add_paragraph()

                if has_arabic(line_text):
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                run = p.add_run(line_text)
                run.font.name = "Arial"

    doc.save(output_path)
