#!/usr/bin/env python3
"""
pdf_to_docx.py — PDF -> DOCX converter with automatic Arabic/RTL support.

WHY THIS EXISTS
----------------
Some PDFs embed fonts that have no ToUnicode CMap (visible as `uni: no` in
`pdffonts`). Those fonts render correctly on screen (the real glyph outlines
are embedded) but there is NO data in the file mapping glyph -> Unicode
character. Any text-layer extraction from such a font produces garbage,
no matter which library you use. This is common in Arabic PDFs generated
by pipelines that fake RTL rendering via glyph substitution on WinAnsi
code points.

STRATEGY
--------
For each page:
  1. Ask `pdffonts` (restricted to that page) which fonts are used.
  2. If any font on that page is embedded + subsetted + has no ToUnicode
     map (and isn't a decorative/bullet font like Wingdings/Symbol), the
     text layer for that page is untrustworthy -> OCR that page instead.
  3. Otherwise, extract the text layer directly (fast, exact).
  4. Either way, detect Arabic-script paragraphs and mark them RTL in the
     output docx (`w:bidi` / `w:rtl`), so Word renders them correctly —
     getting the right Unicode characters is necessary but not sufficient;
     Word also needs to be told which paragraphs are RTL.
  5. Embedded raster images are extracted directly from the PDF (not
     re-rasterized) and inserted at the correct point in the page flow.

REQUIRES (system packages):
  poppler-utils   (pdffonts, pdftotext, pdfimages, pdftoppm / pdftocairo)
  tesseract-ocr
  tesseract-ocr-ara   <-- REQUIRED for Arabic OCR fallback to work.
                          Without it, OCR fallback will silently produce
                          no/garbled Arabic. Install with e.g.:
                          apt-get install tesseract-ocr-ara

REQUIRES (python packages): python-docx, pytesseract, pillow, pdf2image
"""

import argparse
import re
import subprocess
import sys
import tempfile
import unicodedata
from pathlib import Path

from PIL import Image
from pdf2image import convert_from_path
import pytesseract

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------------------------------------------------------------------
# Font risk detection
# ---------------------------------------------------------------------------

DECORATIVE_FONT_HINTS = ("wingdings", "symbol", "webdings", "marlett")


def _pdffonts_table(pdf_path: str, first_page: int = None, last_page: int = None):
    """Run pdffonts (optionally restricted to a page range) and parse rows.

    pdffonts prints a fixed-width table. We use the dashed separator line
    (row 2) to find exact column boundaries instead of guessing from
    whitespace-split tokens, since font names and the "object ID" field
    can each contain internal spaces/multiple tokens.
    """
    cmd = ["pdffonts"]
    if first_page is not None:
        cmd += ["-f", str(first_page)]
    if last_page is not None:
        cmd += ["-l", str(last_page)]
    cmd.append(pdf_path)

    out = subprocess.run(cmd, capture_output=True, text=True, check=True).stdout
    lines = out.splitlines()
    if len(lines) < 2:
        return []

    sep_line = lines[1]  # the '----- ----- -----' line defines column widths
    bounds = []
    start = 0
    for m in re.finditer(r"-+", sep_line):
        bounds.append((start if bounds else m.start(), m.end()))
        start = m.end()
    # bounds now = [(col_start, col_end), ...] for each dashed run, in order:
    # name, type, encoding, emb, sub, uni, object ID

    rows = []
    for line in lines[2:]:
        if not line.strip():
            continue
        cells = []
        for i, (s, e) in enumerate(bounds):
            # last column can overflow the dashed width; take to end of line
            cell = line[s:] if i == len(bounds) - 1 else line[s:e]
            cells.append(cell.strip())
        if len(cells) < 7:
            continue
        name, ftype, encoding, emb, sub, uni, obj_id = cells[:7]
        rows.append({"raw": line, "name": name, "type": ftype,
                      "encoding": encoding, "emb": emb, "sub": sub, "uni": uni})
    return rows


def page_is_font_risky(pdf_path: str, page_num: int) -> bool:
    """True if `page_num` (1-indexed) uses a font whose text layer can't be trusted."""
    try:
        rows = _pdffonts_table(pdf_path, page_num, page_num)
    except subprocess.CalledProcessError:
        # If pdffonts fails for any reason, fail safe -> OCR the page.
        return True

    for row in rows:
        name_lower = row["name"].lower()
        if any(hint in name_lower for hint in DECORATIVE_FONT_HINTS):
            continue  # bullets/icons don't carry body text
        if row["emb"] == "yes" and row["sub"] == "yes" and row["uni"] == "no":
            return True
    return False


# ---------------------------------------------------------------------------
# Text-layer path (fast, used when the page's fonts are trustworthy)
# ---------------------------------------------------------------------------

def extract_page_text_layer(pdf_path: str, page_num: int) -> str:
    cmd = ["pdftotext", "-layout", "-f", str(page_num), "-l", str(page_num), pdf_path, "-"]
    return subprocess.run(cmd, capture_output=True, text=True, check=True).stdout


# ---------------------------------------------------------------------------
# OCR path (used when the page's fonts are NOT trustworthy)
# ---------------------------------------------------------------------------

def ocr_page_text(pdf_path: str, page_num: int, dpi: int = 300, ocr_lang: str = "ara+eng") -> str:
    images = convert_from_path(pdf_path, dpi=dpi, first_page=page_num, last_page=page_num)
    page_img = images[0]
    # psm 6: assume a single uniform block of text -> keeps paragraph structure
    # reasonably intact for lecture-note style documents.
    text = pytesseract.image_to_string(page_img, lang=ocr_lang, config="--psm 6")
    return text


# ---------------------------------------------------------------------------
# Embedded image extraction (kept as real embedded assets, not re-rasterized)
# ---------------------------------------------------------------------------

def extract_page_images(pdf_path: str, page_num: int, out_dir: Path):
    prefix = out_dir / f"p{page_num}_img"
    cmd = ["pdfimages", "-all", "-f", str(page_num), "-l", str(page_num), pdf_path, str(prefix)]
    subprocess.run(cmd, capture_output=True, check=True)
    return sorted(out_dir.glob(f"p{page_num}_img*"))


# ---------------------------------------------------------------------------
# Script detection (per paragraph) — decides RTL vs LTR in the output docx
# ---------------------------------------------------------------------------

ARABIC_RANGES = (
    (0x0600, 0x06FF), (0x0750, 0x077F), (0x08A0, 0x08FF),
    (0xFB50, 0xFDFF), (0xFE70, 0xFEFF),
)


def is_arabic_char(ch: str) -> bool:
    cp = ord(ch)
    return any(lo <= cp <= hi for lo, hi in ARABIC_RANGES)


def paragraph_is_rtl(text: str) -> bool:
    letters = [c for c in text if c.isalpha()]
    if not letters:
        return False
    arabic_count = sum(1 for c in letters if is_arabic_char(c))
    return arabic_count / len(letters) > 0.3  # mixed lines with some Arabic still go RTL


def set_paragraph_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    pPr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in paragraph.runs:
        rPr = run._r.get_or_add_rPr()
        rtl = OxmlElement("w:rtl")
        rPr.append(rtl)


# ---------------------------------------------------------------------------
# Main conversion
# ---------------------------------------------------------------------------

def get_page_count(pdf_path: str) -> int:
    out = subprocess.run(["pdfinfo", pdf_path], capture_output=True, text=True, check=True).stdout
    m = re.search(r"Pages:\s+(\d+)", out)
    return int(m.group(1))


def convert_pdf_to_docx(pdf_path: str, out_path: str, ocr_lang: str = "ara+eng", dpi: int = 300):
    pdf_path = str(pdf_path)
    n_pages = get_page_count(pdf_path)
    doc = Document()

    with tempfile.TemporaryDirectory() as tmp:
        tmp_dir = Path(tmp)

        for page_num in range(1, n_pages + 1):
            risky = page_is_font_risky(pdf_path, page_num)

            if risky:
                raw_text = ocr_page_text(pdf_path, page_num, dpi=dpi, ocr_lang=ocr_lang)
                source = "ocr"
            else:
                raw_text = extract_page_text_layer(pdf_path, page_num)
                source = "text-layer"

            print(f"[page {page_num}] source={source}", file=sys.stderr)

            # Write paragraphs, splitting on blank lines.
            for block in re.split(r"\n\s*\n", raw_text.strip()):
                block = block.strip()
                if not block:
                    continue
                line = " ".join(block.split("\n"))  # join wrapped lines
                p = doc.add_paragraph()
                run = p.add_run(line)
                if paragraph_is_rtl(line):
                    set_paragraph_rtl(p)

            # Insert any embedded images from this page after its text.
            for img_path in extract_page_images(pdf_path, page_num, tmp_dir):
                try:
                    with Image.open(img_path) as im:
                        w_px, h_px = im.size
                    max_width_in = 6.0
                    width_in = min(max_width_in, w_px / 96)
                    height_in = width_in * (h_px / w_px)
                    doc.add_picture(str(img_path), width=Inches(width_in))
                except Exception as e:
                    print(f"  [warn] could not embed image {img_path}: {e}", file=sys.stderr)

            if page_num != n_pages:
                doc.add_page_break()

    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    ap = argparse.ArgumentParser(description="Convert PDF to DOCX with automatic Arabic/RTL + OCR fallback")
    ap.add_argument("pdf", help="input PDF path")
    ap.add_argument("docx", help="output DOCX path")
    ap.add_argument("--lang", default="ara+eng", help="tesseract language(s) for OCR fallback")
    ap.add_argument("--dpi", type=int, default=300, help="OCR render DPI")
    args = ap.parse_args()

    convert_pdf_to_docx(args.pdf, args.docx, ocr_lang=args.lang, dpi=args.dpi)
    print(f"Wrote {args.docx}")
