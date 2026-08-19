import re
import io
import pymupdf as fitz
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

def has_arabic(text: str) -> bool:
    return any('\u0600' <= c <= '\u06FF' for c in text)

def set_rtl_run(run):
    rPr = run._r.get_or_add_rPr()
    rtl = OxmlElement('w:rtl')
    rtl.set(qn('w:val'), '1')
    rPr.append(rtl)

def extract_underlines(page):
    underlines = []
    try:
        drawings = page.get_drawings()
        for d in drawings:
            for item in d.get("items", []):
                if item[0] == "l":
                    p1, p2 = item[1], item[2]
                    # Check if it's horizontal
                    if abs(p1.y - p2.y) < 2:
                        underlines.append({
                            "x0": min(p1.x, p2.x),
                            "x1": max(p1.x, p2.x),
                            "y": p1.y
                        })
    except Exception as e:
        print(f"Warning: could not extract drawings for underlines - {e}")
    return underlines

def is_underlined(span, underlines):
    span_bbox = span["bbox"]
    x0, y0, x1, y1 = span_bbox
    # span y1 is the bottom of the bounding box.
    # An underline is usually slightly below y1 or near the baseline.
    for ul in underlines:
        # Check if the underline overlaps horizontally
        if max(x0, ul["x0"]) < min(x1, ul["x1"]):
            # Check if it's just below the text (within 5 points)
            if 0 <= (ul["y"] - y1) < 5 or abs(ul["y"] - span["origin"][1]) < 3:
                return True
    return False

def convert_pdf_to_word_local(input_pdf_path: str, output_docx_path: str):
    doc_word = docx.Document()
    doc_pdf = fitz.open(input_pdf_path)
    
    for page_num in range(len(doc_pdf)):
        page = doc_pdf[page_num]
        
        # Insert page break if not first page
        if page_num > 0:
            doc_word.add_page_break()
            
        underlines = extract_underlines(page)
        
        blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_IMAGES)["blocks"]
        # Sort blocks vertically, then horizontally to maintain text order perfectly
        blocks.sort(key=lambda b: (b['bbox'][1], b['bbox'][0]))
        
        for block in blocks:
            if block["type"] == 0: # Text block (Paragraph)
                # Group lines into a paragraph
                lines = block["lines"]
                # Sort lines vertically within block just to be safe
                lines.sort(key=lambda l: l['bbox'][1])
                
                # Check if block has Arabic to set paragraph alignment
                block_text = "".join([span["text"] for line in lines for span in line["spans"]])
                is_arabic_block = has_arabic(block_text)
                
                p = doc_word.add_paragraph()
                if is_arabic_block:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                for line in lines:
                    spans = line["spans"]
                    # Sort spans horizontally
                    # If Arabic, PyMuPDF usually returns them in logical order.
                    # Sorting by x0 might break logical order for Arabic, so we trust logical order.
                    if not is_arabic_block:
                        spans.sort(key=lambda s: s['bbox'][0])
                        
                    for span in spans:
                        text = span["text"]
                        if not text.strip():
                            # just a space, add it without complex styling
                            p.add_run(text)
                            continue
                            
                        run = p.add_run(text)
                        
                        # Font Size
                        run.font.size = Pt(span["size"])
                        
                        # Font Style (Bold / Italic)
                        flags = span["flags"]
                        is_bold = flags & 2**4  # PyMuPDF bold flag heuristic
                        is_italic = flags & 2**1 # PyMuPDF italic flag heuristic
                        if "bold" in span["font"].lower():
                            is_bold = True
                        if "italic" in span["font"].lower() or "oblique" in span["font"].lower():
                            is_italic = True
                            
                        run.bold = bool(is_bold)
                        run.italic = bool(is_italic)
                        
                        # Color Identifying
                        color = span["color"]
                        if color != 0: # 0 is black, skip to save space unless explicitly black
                            try:
                                # PyMuPDF color is an integer representing RGB (sRGB)
                                # Extract R, G, B components
                                b = color & 255
                                g = (color >> 8) & 255
                                r = (color >> 16) & 255
                                run.font.color.rgb = RGBColor(r, g, b)
                            except Exception:
                                pass
                                
                        # Underline Identifying
                        if is_underlined(span, underlines):
                            run.underline = True
                            
                        # Arabic Text Processing
                        if is_arabic_block and has_arabic(text):
                            set_rtl_run(run)
                            
            elif block["type"] == 1: # Image block
                image_bytes = block.get("image")
                if image_bytes:
                    image_stream = io.BytesIO(image_bytes)
                    try:
                        # Attempt to add image; Word accepts most PNG/JPEG
                        doc_word.add_picture(image_stream, width=docx.shared.Inches(6.0))
                    except Exception as e:
                        print(f"Warning: could not insert image - {e}")

    doc_pdf.close()
    doc_word.save(output_docx_path)
